#!/usr/bin/env python3
"""
generate_documentation.py
=========================
Reads the CSVs produced by pbi_usage_analyzer.py and generates a table-first
Word document showing full data lineage:

    table -> pages using it -> partitions -> source (server / database / object)
    calculations -> tables they depend on -> those tables' sources

Usage:
    pip install python-docx
    python generate_documentation.py --csv-dir pbi_docs -o pbi_docs/Model_Documentation.docx --title "Sales Model"
"""

import argparse
import csv
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

MUTED = RGBColor(0x60, 0x60, 0x60)


def read_csv(path: Path):
    if not path.exists():
        return []
    with open(path, encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def kv(doc, label, value, indent=0.25):
    """Plain 'Label: value' paragraph — deliberately unstyled."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(indent * 72)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(value if value else "-")
    return p


def muted(doc, text, indent=0.25):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(indent * 72)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.color.rgb = MUTED
    r.font.size = Pt(9)
    return p


def source_string(prow):
    """One-line source description for a partition row."""
    kind = prow["source_kind"]
    if kind.startswith("Embedded binary"):
        return "Embedded binary (enter-data) - opaque payload stored in the model"
    if kind.startswith("Calculated table"):
        return "Calculated table (DAX-derived, no external source)"
    bits = []
    if prow["server"]:
        bits.append(f"server: {prow['server']}")
    if prow["database"]:
        bits.append(f"database: {prow['database']}")
    if prow["source_objects"]:
        bits.append(f"objects: {prow['source_objects']}")
    return f"{kind}" + (f" ({'; '.join(bits)})" if bits else "")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--csv-dir", required=True, help="Folder containing the analyzer CSVs")
    ap.add_argument("-o", "--out", default=None, help="Output .docx path")
    ap.add_argument("--title", default="Power BI Model", help="Model/report name for the title")
    args = ap.parse_args()

    d = Path(args.csv_dir)
    out_path = Path(args.out) if args.out else d / "Model_Documentation.docx"

    summary = read_csv(d / "table_usage_summary.csv")
    partitions = read_csv(d / "partitions.csv")
    measures = read_csv(d / "measures.csv")
    calc_cols = read_csv(d / "calculated_columns.csv")
    relationships = read_csv(d / "relationships.csv")

    # ---- indexes -------------------------------------------------------------
    parts_by_table = defaultdict(dict)       # table -> {partition_name: row}
    pages_by_table = defaultdict(dict)       # table -> {page: usage_type}
    for p in partitions:
        parts_by_table[p["table"]].setdefault(p["partition"], p)
        if p["page"]:
            pages_by_table[p["table"]][p["page"]] = p["usage_type"]

    measures_by_table = defaultdict(list)
    for m in measures:
        measures_by_table[m["home_table"]].append(m)
    calcs_by_table = defaultdict(list)
    for c in calc_cols:
        calcs_by_table[c["table"]].append(c)

    def table_source_summary(tname):
        """Compact 'where does this table come from' line, for cross-references."""
        rows = list(parts_by_table.get(tname, {}).values())
        if not rows:
            return "no partition information"
        return " | ".join(sorted({source_string(r) for r in rows}))

    model_tables = [s for s in summary if "NOT IN MODEL" not in s["classification"]]
    phantom = [s for s in summary if "NOT IN MODEL" in s["classification"]]
    unused = [s for s in model_tables if s["classification"] == "UNUSED"]

    # ---- document ------------------------------------------------------------
    doc = Document()
    doc.add_heading(f"{args.title} - Data Lineage Documentation", level=0)
    doc.add_paragraph(f"Generated {date.today().isoformat()} from static analysis "
                      f"of the semantic model and PBIR report definitions.")

    # 1. Overview
    doc.add_heading("1. Overview", level=1)
    kv(doc, "Tables in model", str(len(model_tables)), indent=0)
    kv(doc, "Tables used directly on pages",
       str(sum(1 for s in model_tables if s["classification"] == "Used - direct")), indent=0)
    kv(doc, "Tables used only via measure DAX",
       str(sum(1 for s in model_tables if s["classification"] == "Used - via measure DAX")), indent=0)
    kv(doc, "Tables possibly used (relationship only)",
       str(sum(1 for s in model_tables if s["classification"].startswith("Possibly"))), indent=0)
    kv(doc, "Unused tables", "; ".join(s["table"] for s in unused) or "none", indent=0)
    if phantom:
        kv(doc, "Referenced by reports but NOT in model",
           "; ".join(s["table"] for s in phantom), indent=0)
        muted(doc, "These references indicate a renamed table or a report pointing "
                   "at a different semantic model. See Appendix A.", indent=0)

    # 2. Tables
    doc.add_heading("2. Tables", level=1)
    for s in model_tables:
        tname = s["table"]
        doc.add_heading(tname, level=2)

        flags = [s["classification"]]
        if s["is_hidden"] == "True":
            flags.append("hidden")
        if s["is_calculated_table"] == "True":
            flags.append("calculated table")
        kv(doc, "Status", " | ".join(flags), indent=0)

        # pages
        pages = pages_by_table.get(tname, {})
        if pages:
            kv(doc, "Used on pages", "", indent=0)
            for page, usage in sorted(pages.items()):
                muted(doc, f"- {page}  ({usage})")
        else:
            kv(doc, "Used on pages", "none", indent=0)

        # partitions / sources
        kv(doc, "Partitions (data sources)", "", indent=0)
        prows = parts_by_table.get(tname, {})
        if not prows:
            muted(doc, "- no partitions found in model")
        for pname, prow in sorted(prows.items()):
            muted(doc, f"- {pname}  [{prow['mode']}]  ->  {source_string(prow)}")
            if prow["note"] and not prow["source_kind"].startswith("Embedded binary"):
                muted(doc, f"    note: {prow['note']}", indent=0.45)

        # measures homed here
        if measures_by_table.get(tname):
            kv(doc, "Measures in this table", "", indent=0)
            for m in sorted(measures_by_table[tname], key=lambda x: x["measure"]):
                used = f"used on: {m['pages']}" if m["pages"] else "not used on any page"
                muted(doc, f"- [{m['measure']}]  ({used})")
                deps = [t.strip() for t in m["tables_referenced_in_dax"].split(";") if t.strip()]
                for dep in deps:
                    muted(doc, f"    depends on {dep}  ->  {table_source_summary(dep)}",
                          indent=0.45)

        # calculated columns homed here
        tcalcs = [c for c in calcs_by_table.get(tname, []) if c["object_type"] == "calculated column"]
        if tcalcs:
            kv(doc, "Calculated columns", "", indent=0)
            for c in sorted(tcalcs, key=lambda x: x["name"]):
                refs = [t.strip() for t in c["tables_referenced"].split(";") if t.strip()]
                muted(doc, f"- {c['name']}")
                for dep in refs:
                    muted(doc, f"    depends on {dep}  ->  {table_source_summary(dep)}",
                          indent=0.45)

    # 3. Relationships
    if relationships:
        doc.add_heading("3. Relationships", level=1)
        for r in relationships:
            active = "active" if r["is_active"] == "True" else "INACTIVE"
            muted(doc, f"- {r['from_table']}[{r['from_column']}] -> "
                       f"{r['to_table']}[{r['to_column']}]  "
                       f"({active}, {r['cross_filtering']})", indent=0)

    # Appendix A: phantom references
    if phantom:
        doc.add_heading("Appendix A - Report references not found in the model", level=1)
        for s in phantom:
            muted(doc, f"- {s['table']}  (referenced on: {s['pages_direct']})", indent=0)

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
